import io
import os
import re
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from pydub import AudioSegment
from openai import OpenAI
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load environment variables from .env file
load_dotenv()

# Instantiate the OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def convert_to_mp3(file, filename):
    ext = filename.split('.')[-1].lower()
    print(f"File extension detected: {ext}")

    if (ext == 'mp3'):
        print("File is already in MP3 format, no conversion needed.")
        file.seek(0)
        mp3_buffer = io.BytesIO(file.read())
        mp3_buffer.name = filename
        return mp3_buffer
    else:
        print(f"Converting {ext} to MP3...")
        audio = AudioSegment.from_file(file, format=ext)
        mp3_buffer = io.BytesIO()
        audio.export(mp3_buffer, format="mp3")
        mp3_buffer.name = "audio.mp3"
        mp3_buffer.seek(0)
        return mp3_buffer

def add_paragraph_with_direction(doc, text):
    # Detect if text is RTL or LTR based on the presence of Hebrew/Arabic characters
    rtl_pattern = re.compile(r'[\u0590-\u05FF\u0600-\u06FF]')
    is_rtl = rtl_pattern.search(text)

    paragraph = doc.add_paragraph(text.strip())

    if is_rtl:
        paragraph.alignment = 2  # Right alignment
        set_rtl(paragraph)
    else:
        paragraph.alignment = 0  # Left alignment

def set_rtl(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    # Set RTL direction
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), "1")
    pPr.append(bidi)

    # Explicitly set paragraph alignment to right
    alignment = OxmlElement('w:jc')
    alignment.set(qn('w:val'), "right")
    pPr.append(alignment)

def generate_srt(segments, is_translation=False):
    srt_content = ""
    for i, segment in enumerate(segments, start=1):
        if not is_translation and 'start' in segment and 'end' in segment:
            start_time = format_time(segment['start'])
            end_time = format_time(segment['end'])
        else:
            start_time = format_time(i * 5)
            end_time = format_time((i + 1) * 5)
        text = segment['text']
        srt_content += f"{i}\n{start_time} --> {end_time}\n{text}\n\n"
    return srt_content

def generate_vtt(segments, is_translation=False):
    vtt_content = "WEBVTT\n\n"
    for i, segment in enumerate(segments, start=1):
        if not is_translation and 'start' in segment and 'end' in segment:
            start_time = format_time(segment['start']).replace(',', '.')
            end_time = format_time(segment['end']).replace(',', '.')
        else:
            start_time = format_time(i * 5).replace(',', '.')
            end_time = format_time((i + 1) * 5).replace(',', '.')
        text = segment['text']
        vtt_content += f"{start_time} --> {end_time}\n{text}\n\n"
    return vtt_content

def format_time(seconds):
    if isinstance(seconds, str):
        try:
            seconds = float(seconds)
        except ValueError:
            seconds = 0
    
    hours = int(seconds / 3600)
    minutes = int((seconds % 3600) / 60)
    seconds = seconds % 60
    milliseconds = int((seconds - int(seconds)) * 1000)
    return f"{hours:02d}:{minutes:02d}:{int(seconds):02d},{milliseconds:03d}"

def create_app():
    app = Flask(__name__)
    CORS(app)

    # Enable template auto-reloading
    app.config['TEMPLATES_AUTO_RELOAD'] = True

    @app.route("/")
    def index():
        return render_template("index.html")

    @app.route("/transcribe", methods=["POST"])
    def transcribe():
        try:
            file = request.files["audio"]
            filename = file.filename
            print(f"Received file: {filename}")

            # Convert the uploaded file to mp3 if necessary
            mp3_buffer = convert_to_mp3(file, filename)
            print("File conversion successful.")

            # Request transcription with timecodes
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=mp3_buffer,
                response_format="verbose_json"  # Get detailed JSON response including timecodes
            )
            print("Transcription successful.")

            # Ensure each segment ends with a period
            for segment in transcript.segments:
                if not segment['text'].strip().endswith('.'):
                    segment['text'] = segment['text'].strip() + '.'

            return jsonify({"output": transcript.segments})

        except Exception as e:
            print(f"Error during transcription: {e}")
            return jsonify({"error": str(e)}), 500

    @app.route("/translate", methods=["POST"])
    def translate():
        try:
            data = request.get_json()
            segments = data.get('segments', [])
            target_language = data.get('language', '')

            print(f"Received segments for translation: {len(segments)} segments")
            print(f"Target language: {target_language}")

            if not segments or not target_language:
                raise ValueError("Segments or language not provided")

            translated_segments = []
            for segment in segments:
                text = segment['text']
                prompt = f"Translate the following text to {target_language}:\n\n{text}\n\nOutput only the translated text."

                translation = client.chat.completions.create(
                    model="gpt-3.5-turbo",  # Changed from "gpt-4o-mini" to "gpt-3.5-turbo"
                    messages=[{"role": "user", "content": prompt}]
                )

                translated_text = translation.choices[0].message.content
                translated_segments.append({
                    'start': segment.get('start'),
                    'end': segment.get('end'),
                    'text': translated_text
                })

            print(f"Translation successful: {len(translated_segments)} segments translated")
            return jsonify({"output": translated_segments})

        except ValueError as ve:
            print(f"Value error: {ve}")
            return jsonify({"error": str(ve)}), 400
        except Exception as e:
            print(f"Error during translation: {e}")
            return jsonify({"error": "An unknown error occurred"}), 500

    @app.route("/download", methods=["POST"])
    def download():
        try:
            data = request.get_json()
            print("Received data:", data)  # Log the received data
            segments = data.get('segments', [])
            file_format = data.get('format', '')
            download_type = data.get('type', 'transcription')

            if not segments or not file_format:
                print("Missing segments or format")  # Log if segments or format is missing
                return jsonify({"error": "Segments or format missing from request"}), 400

            is_translation = download_type == 'translation'

            if file_format == 'docx':
                try:
                    doc = Document()
                    for segment in segments:
                        if 'start' in segment and 'end' in segment:
                            timestamp = f"[{format_time(segment['start'])} --> {format_time(segment['end'])}] "
                        else:
                            timestamp = ""
                        text = segment.get('text', '') if isinstance(segment, dict) else segment
                        add_paragraph_with_direction(doc, timestamp + text)
                    
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return send_file(
                        buffer, 
                        as_attachment=True, 
                        download_name=f"{download_type}.docx", 
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    app.logger.error(f"Error generating DOCX: {e}")
                    return jsonify({"error": "Error generating DOCX"}), 500

            elif file_format == 'srt':
                try:
                    srt_content = generate_srt(segments, is_translation)
                    return send_file(
                        io.BytesIO(srt_content.encode('utf-8')),
                        as_attachment=True,
                        download_name=f"{download_type}.srt",
                        mimetype="text/plain"
                    )
                except Exception as e:
                    app.logger.error(f"Error generating SRT: {e}")
                    return jsonify({"error": f"Error generating SRT: {str(e)}"}), 500

            elif file_format == 'vtt':
                try:
                    vtt_content = generate_vtt(segments, is_translation)
                    return send_file(
                        io.BytesIO(vtt_content.encode('utf-8')),
                        as_attachment=True,
                        download_name=f"{download_type}.vtt",
                        mimetype="text/vtt"
                    )
                except Exception as e:
                    app.logger.error(f"Error generating VTT: {e}")
                    return jsonify({"error": f"Error generating VTT: {str(e)}"}), 500

            else:
                return jsonify({"error": "Unsupported format"}), 400

        except Exception as e:
            app.logger.error(f"Error in download endpoint: {e}")
            return jsonify({"error": f"Internal Server Error: {str(e)}"}), 500

    return app

# Create the Flask app instance at the module level
app = create_app()

if __name__ == "__main__":
    app.run(debug=True)